from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from estimator import JobInput, EstimateResult


def generate_contract_docx(job: JobInput, estimate: EstimateResult, selling_price: float) -> Document:
    doc = Document()

    # Header block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Miller Roofing NJ LLC\n")
    run.bold = True
    run.font.size = Pt(24)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("PO Box 221 Succasunna, NJ 07836 (mailing) • 221 Tucker Drive, Unit 14 (physical)\n")
    p2.add_run("Phone: (973) 555-5555 • Fax: (973) 675-4321\n")
    p2.add_run(f"Email: {job.salesperson_email}")

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p3.add_run("NJ Home Improvement Contractor License — 28ES15131902")
    run.bold = True

    doc.add_paragraph()

    # Date
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.add_run(date.today().strftime("%B %d, %Y"))

    # Customer block
    doc.add_paragraph(job.customer_name)
    doc.add_paragraph(job.address)
    if job.customer_phone:
        doc.add_paragraph(job.customer_phone)
    if job.customer_email:
        doc.add_paragraph(job.customer_email)

    doc.add_paragraph()
    doc.add_paragraph(f"Dear {job.customer_name.split()[0]},")

    doc.add_paragraph(
        "Thank you for meeting with me to discuss your home improvement project. "
        "Included in this proposal is the estimate for the work we discussed, including "
        "a complete replacement of your roof."
    )

    doc.add_paragraph(
        "Please visit our website, www.milleroofingnj.com, or our Facebook page for more information about our company and services."
    )

    # Salesperson block
    doc.add_heading("Prepared By", level=1)
    doc.add_paragraph(f"Salesperson: {job.salesperson_name}")
    doc.add_paragraph(f"Phone: {job.salesperson_phone}")
    doc.add_paragraph(f"Email: {job.salesperson_email}")

    # Project summary
    doc.add_heading("Project Summary", level=1)
    doc.add_paragraph(f"Selected Package: {estimate.package_name}")
    doc.add_paragraph(f"Shingle System: {estimate.shingle_name}")
    doc.add_paragraph(f"Warranty: {estimate.warranty_name}")
    doc.add_paragraph(f"Total Squares: {estimate.total_squares}")
    doc.add_paragraph(f"Estimated Selling Price: ${selling_price:,.2f}")

    # Roof details
    doc.add_heading("Roof Measurements", level=1)
    doc.add_paragraph(f"Squares (0-7 pitch): {job.squares_0_7}")
    doc.add_paragraph(f"Squares (8-11 pitch): {job.squares_8_11}")
    doc.add_paragraph(f"Squares (12-16 pitch): {job.squares_12_16}")
    doc.add_paragraph(f"Eaves: {job.eaves_feet} ft")
    doc.add_paragraph(f"Rakes: {job.rakes_feet} ft")
    doc.add_paragraph(f"Ridge Vent: {job.ridge_vent_feet} ft")

    # Estimate breakdown
    doc.add_heading("Estimate Breakdown", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Qty"
    hdr[2].text = "Unit Price"
    hdr[3].text = "Total"

    for line in estimate.lines:
        row = table.add_row().cells
        row[0].text = line.name
        row[1].text = str(line.quantity)
        row[2].text = f"${line.unit_price:,.2f}"
        row[3].text = f"${line.total:,.2f}"

    doc.add_heading("Total Project Price", level=1)
    doc.add_paragraph(f"${selling_price:,.2f}")

    doc.add_heading("Scope of Work", level=1)
    doc.add_paragraph(
        "Remove existing roofing materials and install a complete roofing system including shingles, "
        "underlayment, ventilation, flashing, and all listed components. Work includes cleanup and proper disposal."
    )

    doc.add_paragraph(
        "Pricing is calculated using a structured estimation system that accounts for roof pitch, "
        "material selection, and accessory requirements to improve consistency and accuracy."
    )

    doc.add_heading("Payment Terms", level=1)
    doc.add_paragraph("Deposit due upon signing. Balance due upon completion.")

    doc.add_heading("Signatures", level=1)
    doc.add_paragraph("Homeowner: ________________________    Date: __________")
    doc.add_paragraph("Miller Roofing NJ Rep: ________________________    Date: __________")

    return doc


def save_contract(job: JobInput, estimate: EstimateResult, selling_price: float, output_dir: str = "outputs") -> str:
    Path(output_dir).mkdir(parents=True, exist_ok=True)

    safe_name = job.customer_name.lower().replace(" ", "_")
    safe_package = estimate.package_name.lower()
    filename = f"{safe_name}_{safe_package}_proposal.docx"
    file_path = Path(output_dir) / filename

    doc = generate_contract_docx(job, estimate, selling_price)
    doc.save(file_path)

    return str(file_path)